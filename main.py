# ==========================================
# MindMorph - Pro Subliminal Audio Editor
# Version: 4.9 (OOP - Project Save/Load Implemented)
# ==========================================
# --- Early Config ---
import os

import streamlit as st
from PIL import Image

favicon_path = os.path.join("assets", "favico.png")
page_icon = None
try:
    page_icon = Image.open(favicon_path)
except FileNotFoundError:
    pass
st.set_page_config(layout="wide", page_title="MindMorph - Pro Subliminal Editor", page_icon=page_icon)

# --- Imports ---
import hashlib  # For creating setting hash
import json  # For project save/load
import logging
import logging.handlers
import math  # For fading (No longer used, but keep import for now)
import queue
import re  # For sanitizing filename
import tempfile
import textwrap
import time
import traceback
import uuid
from io import BytesIO, StringIO
from typing import Any, Dict, List, Optional, Tuple

# --- Logging Setup ---
# (Same setup as before - creates editor_oop.log)
log_queue = queue.Queue(-1)
log_file = "editor_oop.log"
log_formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s - [%(funcName)s] - %(message)s")
file_handler = logging.handlers.RotatingFileHandler(log_file, maxBytes=10 * 1024 * 1024, backupCount=3)  # 10MB
file_handler.setFormatter(log_formatter)
file_handler.setLevel(logging.DEBUG)
queue_handler = logging.handlers.QueueHandler(log_queue)
root_logger = logging.getLogger()
root_logger.setLevel(logging.DEBUG)
if not root_logger.hasHandlers() or not any(isinstance(h, logging.handlers.QueueHandler) for h in root_logger.handlers):
    root_logger.addHandler(queue_handler)
if "listener_started" not in st.session_state:
    listener = logging.handlers.QueueListener(log_queue, file_handler, respect_handler_level=True)
    listener.start()
    st.session_state.listener_started = True
logger = logging.getLogger(__name__)
logger.info("-----------------------------------------------------")
logger.info("Application starting up / rerunning.")
logger.info("-----------------------------------------------------")

# --- Core Libraries ---
try:
    import docx  # For reading .docx files
    import librosa
    import librosa.effects
    import numpy as np
    import pyttsx3
    import soundfile as sf
    from scipy import signal
    from streamlit.runtime.uploaded_file_manager import UploadedFile
    from streamlit_advanced_audio import WaveSurferOptions, audix

    # Try importing pydub for MP3 export check
    try:
        from pydub import AudioSegment

        PYDUB_AVAILABLE = True
        logger.info("pydub library found. MP3 export might be available (if ffmpeg is also installed).")
    except ImportError:
        PYDUB_AVAILABLE = False
        logger.warning("pydub library not found. MP3 export will be disabled.")

    logger.debug("Core audio, UI, and docx libraries imported successfully.")
except ImportError as e:
    logger.exception("CRITICAL: Failed to import core libraries. App cannot continue.")
    missing_lib = e.name
    install_cmd = f"pip install {missing_lib}"
    if missing_lib == "docx":
        install_cmd = "pip install python-docx"
    elif missing_lib == "pydub":
        install_cmd = "pip install pydub"
    st.error(f"Core library import failed: {e}. Ensure dependencies installed.")
    st.error(f"Missing: {missing_lib}")
    st.error(f"Try: `{install_cmd}`")
    st.code(traceback.format_exc())
    st.stop()

# --- Constants ---
GLOBAL_SR = 44100
logger.debug(f"Global Sample Rate set to: {GLOBAL_SR} Hz")
TTS_CHUNK_SIZE = 1500  # For full generation chunking
PREVIEW_DURATION_S = 60  # Default duration for previews in editor
MIX_PREVIEW_DURATION_S = 10  # Duration for the master mix preview
MIX_PREVIEW_PROCESSING_BUFFER_S = 5  # Extra seconds to process for preview to handle speed changes
logger.debug(f"Editor Preview duration: {PREVIEW_DURATION_S}s")
logger.debug(f"Mix Preview duration: {MIX_PREVIEW_DURATION_S}s")
PROJECT_FILE_VERSION = "1.0"  # For future compatibility checks

# --- Data Types ---
AudioData = np.ndarray
SampleRate = int
TrackID = str
TrackData = Dict[str, Any]
TrackType = str  # e.g., 'affirmation', 'background', 'frequency', 'voice', 'other'

# --- Constants for Track Types ---
TRACK_TYPE_AFFIRMATION = "🗣️ Affirmation"
TRACK_TYPE_BACKGROUND = "🎵 Background/Mask"
TRACK_TYPE_FREQUENCY = "🧠 Frequency"
TRACK_TYPE_VOICE = "🎤 Voice Recording"
TRACK_TYPE_OTHER = "⚪ Other"
TRACK_TYPES = [TRACK_TYPE_AFFIRMATION, TRACK_TYPE_BACKGROUND, TRACK_TYPE_FREQUENCY, TRACK_TYPE_VOICE, TRACK_TYPE_OTHER]


# ==========================================
# 1. Audio Processing Utilities (Functions)
# ==========================================
# (Functions load_audio, save_audio, save_audio_to_temp, generate*, apply*, etc. remain unchanged from v4.0)
# ... (Keep all audio utility functions from the previous version here) ...
def load_audio(file_source: UploadedFile | BytesIO | str, target_sr: SampleRate = GLOBAL_SR) -> tuple[AudioData, SampleRate]:
    """Loads, ensures stereo, and resamples audio."""
    logger.info(f"Loading audio from source type: {type(file_source)}")
    try:
        source = file_source
        audio, sr = librosa.load(source, sr=None, mono=False)
        logger.debug(f"Loaded audio original SR: {sr}, shape: {audio.shape}")
        if audio.ndim == 1:
            audio = np.stack([audio, audio], axis=-1)
        elif audio.shape[0] == 2 and audio.shape[1] > 2:
            audio = audio.T  # Check if channels are first dim and there are samples
        if audio.shape[1] > 2:
            logger.warning(f"Audio > 2 channels ({audio.shape[1]}). Using first two.")
            audio = audio[:, :2]
        elif audio.shape[1] == 1:
            logger.warning("Mono audio loaded unexpectedly. Duplicating.")
            audio = np.concatenate([audio, audio], axis=1)
        if sr != target_sr:
            logger.info(f"Resampling from {sr} Hz to {target_sr} Hz.")
            if audio.size > 0:
                audio = librosa.resample(audio.T, orig_sr=sr, target_sr=target_sr).T
            else:
                sr = target_sr
        return audio.astype(np.float32), target_sr
    except Exception as e:
        logger.exception("Error loading audio.")
        st.error(f"Error loading audio: {e}")
        return np.zeros((0, 2), dtype=np.float32), target_sr


def save_audio(audio: AudioData, sr: SampleRate) -> BytesIO:
    """Saves audio to an in-memory BytesIO buffer (WAV format, PCM16)."""
    buffer = BytesIO()
    logger.debug(f"Saving audio ({audio.shape}, {sr}Hz) to BytesIO.")
    try:
        audio_int16 = (audio * 32767).astype(np.int16)
        sf.write(buffer, audio_int16, sr, format="WAV", subtype="PCM_16")
        buffer.seek(0)
    except Exception as e:
        logger.exception("Error saving audio to buffer.")
        st.error(f"Error saving: {e}")
        buffer = BytesIO()
    return buffer


def save_audio_to_temp(audio: AudioData, sr: SampleRate) -> str | None:
    """Saves audio to a temporary WAV file on disk (WAV format, PCM16). Returns file path or None."""
    temp_file_path = None
    logger.debug(f"Attempting to save audio ({audio.shape}, {sr}Hz) to temp file.")
    try:
        audio_int16 = (audio * 32767).astype(np.int16)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav", mode="wb") as tmp:
            sf.write(tmp, audio_int16, sr, format="WAV", subtype="PCM_16")
            temp_file_path = tmp.name
        logger.info(f"Audio saved successfully to temporary file: {temp_file_path}")
        return temp_file_path
    except Exception as e:
        logger.exception("Failed to save temporary audio file.")
        st.error(f"Failed to save temp file: {e}")
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
                logger.debug(f"Cleaned up partial temp file: {temp_file_path}")
            except OSError as e_os:
                logger.warning(f"Failed cleanup {temp_file_path}: {e_os}")
        return None


def generate_binaural_beats(duration: float, freq_left: float, freq_right: float, sr: SampleRate, volume: float) -> AudioData:
    logger.info(f"Generating binaural beats: dur={duration}s, L={freq_left}Hz, R={freq_right}Hz, vol={volume}")
    t = np.linspace(0, duration, int(sr * duration), endpoint=False)
    left = volume * np.sin(2 * np.pi * freq_left * t)
    right = volume * np.sin(2 * np.pi * freq_right * t)
    audio = np.stack([left, right], axis=1).astype(np.float32)
    return np.clip(audio, -1.0, 1.0)


def generate_solfeggio_frequency(duration: float, freq: float, sr: SampleRate, volume: float) -> AudioData:
    logger.info(f"Generating Solfeggio: dur={duration}s, F={freq}Hz, vol={volume}")
    t = np.linspace(0, duration, int(sr * duration), endpoint=False)
    sine_wave = volume * np.sin(2 * np.pi * freq * t)
    audio = np.stack([sine_wave, sine_wave], axis=1).astype(np.float32)
    return np.clip(audio, -1.0, 1.0)


def generate_isochronic_tones(duration: float, carrier_freq: float, pulse_freq: float, sr: SampleRate, volume: float) -> AudioData:
    """Generates stereo isochronic tones (amplitude modulation)."""
    logger.info(f"Generating Isochronic: dur={duration}s, Carrier={carrier_freq}Hz, Pulse={pulse_freq}Hz, vol={volume}")
    num_samples = int(sr * duration)
    t = np.linspace(0, duration, num_samples, endpoint=False)
    carrier_wave = np.sin(2 * np.pi * carrier_freq * t)
    modulation_wave = 0.5 * (np.sign(np.sin(2 * np.pi * pulse_freq * t + np.pi / 2)) + 1)
    modulated_signal = carrier_wave * modulation_wave
    audio = np.stack([modulated_signal, modulated_signal], axis=1) * volume
    return np.clip(audio, -1.0, 1.0).astype(np.float32)


def generate_white_noise(num_samples: int, sr: SampleRate) -> AudioData:
    """Generates stereo white noise (uniform distribution)."""
    logger.debug(f"Generating {num_samples} samples of white noise.")
    return np.random.uniform(-1.0, 1.0, size=(num_samples, 2)).astype(np.float32)


def generate_pink_noise(num_samples: int, sr: SampleRate) -> AudioData:
    """Generates stereo pink noise (approximated by filtering white noise)."""
    logger.debug(f"Generating {num_samples} samples of pink noise.")
    white = generate_white_noise(num_samples, sr)
    pink = np.cumsum(white, axis=0)
    max_val = np.max(np.abs(pink))
    if max_val > 1e-6:
        pink /= max_val
    return pink.astype(np.float32)


def generate_brown_noise(num_samples: int, sr: SampleRate) -> AudioData:
    """Generates stereo brown noise (approximated by integrating white noise)."""
    logger.debug(f"Generating {num_samples} samples of brown noise.")
    white = generate_white_noise(num_samples, sr)
    brown = np.cumsum(white, axis=0)
    max_val = np.max(np.abs(brown))
    if max_val > 1e-6:
        brown /= max_val
    return brown.astype(np.float32)


def generate_noise(noise_type: str, duration: float, sr: SampleRate, volume: float) -> Optional[AudioData]:
    """Generates specified noise type."""
    logger.info(f"Generating {noise_type}: dur={duration}s, vol={volume}")
    num_samples = int(sr * duration)
    noise_audio = None
    if noise_type == "White Noise":
        noise_audio = generate_white_noise(num_samples, sr)
    elif noise_type == "Pink Noise":
        noise_audio = generate_pink_noise(num_samples, sr)
    elif noise_type == "Brown Noise":
        noise_audio = generate_brown_noise(num_samples, sr)
    else:
        logger.error(f"Unknown noise type requested: {noise_type}")
        return None
    if noise_audio is not None:
        audio = np.clip(noise_audio * volume, -1.0, 1.0)
        return audio.astype(np.float32)
    return None


def apply_reverse(audio: AudioData) -> AudioData:
    """Reverses the audio data along the time axis."""
    logger.debug("Applying reverse effect.")
    if audio is None or audio.size == 0:
        return audio
    try:
        return audio[::-1].astype(np.float32)
    except Exception as e:
        logger.exception("Error applying reverse effect.")
        st.error(f"Reverse effect failed: {e}")
        return audio


def apply_speed_change(audio: AudioData, sr: SampleRate, speed_factor: float) -> AudioData:
    logger.debug(f"Applying speed change factor: {speed_factor} (time_stretch)")
    if np.any(np.isnan(audio)) or np.any(np.isinf(audio)):
        logging.error("NaN/Inf in audio data before time stretch.")
        st.error("Bad audio data.")
        return audio
    if not np.isclose(speed_factor, 1.0):
        if speed_factor <= 0:
            logging.warning(f"Invalid speed factor: {speed_factor}.")
            st.warning("Speed > 0 required.")
            return audio
        try:
            if audio.size == 0:
                return audio
            logging.info(f"Applying time stretch (factor: {speed_factor})...")
            audio_contiguous = np.ascontiguousarray(audio.T)
            audio_stretched = librosa.effects.time_stretch(audio_contiguous, rate=speed_factor)
            return audio_stretched.T.astype(np.float32)
        except Exception as e:
            logging.exception("Error in time_stretch.")
            st.error(f"Speed change failed: {e}")
            return audio
    return audio


def apply_pitch_shift(audio: AudioData, sr: SampleRate, n_steps: float) -> AudioData:
    logger.debug(f"Applying pitch shift: {n_steps} semitones")
    if not np.isclose(n_steps, 0.0):
        try:
            if audio.size == 0:
                return audio
            logging.info(f"Applying pitch shift ({n_steps} semitones)...")
            audio_contiguous = np.ascontiguousarray(audio.T)
            audio_shifted = librosa.effects.pitch_shift(audio_contiguous, sr=sr, n_steps=n_steps)
            return audio_shifted.T.astype(np.float32)
        except Exception as e:
            logging.exception("Error applying pitch shift.")
            st.error(f"Pitch shift failed: {e}")
            return audio
    return audio


def apply_filter(audio: AudioData, sr: SampleRate, filter_type: str, cutoff: float) -> AudioData:
    logger.debug(f"Applying filter: type={filter_type}, cutoff={cutoff} Hz")
    if filter_type == "off" or cutoff <= 0:
        return audio
    try:
        if audio.size == 0:
            return audio
        nyquist = 0.5 * sr
        normalized_cutoff = cutoff / nyquist
        if normalized_cutoff >= 1.0:
            msg = f"Filter cutoff ({cutoff}) >= Nyquist ({nyquist})."
            logging.warning(msg)
            st.warning(msg)
            return audio
        if normalized_cutoff <= 0:
            msg = f"Filter cutoff ({cutoff}) must be positive."
            logging.warning(msg)
            st.warning(msg)
            return audio
        filter_order = 4
        logging.info(f"Applying {filter_type} filter, cutoff {cutoff} Hz, order {filter_order}.")
        if filter_type == "lowpass":
            b, a = signal.butter(filter_order, normalized_cutoff, btype="low")
        elif filter_type == "highpass":
            b, a = signal.butter(filter_order, normalized_cutoff, btype="high")
        else:
            logging.warning(f"Unknown filter type: {filter_type}")
            st.warning(f"Unknown filter: {filter_type}")
            return audio
        min_len_filtfilt = signal.filtfilt_coeffs(b, a)[-1].size * 3
        if audio.shape[0] <= min_len_filtfilt:
            msg = f"Track too short ({audio.shape[0]}) for filter. Skipping."
            logging.warning(msg)
            st.warning(msg)
            return audio
        audio_filtered = signal.filtfilt(b, a, audio, axis=0)
        return audio_filtered.astype(np.float32)
    except Exception as e:
        logging.exception("Error applying filter.")
        st.error(f"Filter failed: {e}")
        return audio


def apply_all_effects(track_data: TrackData, audio_segment: Optional[AudioData] = None) -> AudioData:
    """Applies effects sequentially: Reverse (optional) -> Speed -> Pitch -> Filter."""
    track_name = track_data.get("name", "Unnamed")
    should_reverse = track_data.get("reverse_audio", False)
    if audio_segment is not None:
        audio = audio_segment.copy()
        logger.debug(f"Applying effects to segment for: '{track_name}' (Reverse={should_reverse})")
    else:
        logger.info(f"Applying effects to full original_audio for: '{track_name}' (Reverse={should_reverse})")
        audio = track_data.get("original_audio")
    if audio is None:
        logging.error(f"'{track_name}' missing 'original_audio'.")
        return np.zeros((0, 2), dtype=np.float32)
    audio = audio.copy()
    sr = track_data.get("sr", GLOBAL_SR)
    if audio.size == 0:
        logging.warning(f"Input audio for effects is empty for '{track_name}'.")
        return audio
    if should_reverse:
        audio = apply_reverse(audio)
    audio = apply_speed_change(audio, sr, track_data.get("speed_factor", 1.0))
    audio = apply_pitch_shift(audio, sr, track_data.get("pitch_shift", 0))
    audio = apply_filter(audio, sr, track_data.get("filter_type", "off"), track_data.get("filter_cutoff", 8000.0))
    logger.debug(f"Finished applying base effects for '{track_name}'")
    return audio


def get_preview_audio(track_data: TrackData, preview_duration_s: int = PREVIEW_DURATION_S) -> Optional[AudioData]:
    """Generates a preview (first N seconds) of the track with effects AND volume/pan applied."""
    track_name = track_data.get("name", "N/A")
    logger.info(f"Generating preview audio for track '{track_name}'")
    original_audio = track_data.get("original_audio")
    sr = track_data.get("sr", GLOBAL_SR)
    if original_audio is None or original_audio.size == 0:
        logger.warning(f"No original audio for '{track_name}'.")
        return None
    try:
        preview_samples = min(len(original_audio), int(sr * preview_duration_s))
        if preview_samples <= 0:
            logger.warning(f"Preview samples <= 0 for '{track_name}'.")
            return None
        preview_segment = original_audio[:preview_samples].copy()
        logger.debug(f"Applying effects to preview segment (len={preview_samples}) for '{track_name}'")
        processed_preview = apply_all_effects(track_data, audio_segment=preview_segment)
        vol = track_data.get("volume", 1.0)
        pan = track_data.get("pan", 0.0)
        logger.debug(f"Applying Vol ({vol:.2f}) / Pan ({pan:.2f}) to preview for '{track_name}'")
        pan_rad = (pan + 1) * np.pi / 4
        left_gain = vol * np.cos(pan_rad)
        right_gain = vol * np.sin(pan_rad)
        if processed_preview.ndim == 2 and processed_preview.shape[1] == 2:
            processed_preview[:, 0] *= left_gain
            processed_preview[:, 1] *= right_gain
        else:
            logger.warning(f"Preview for '{track_name}' not stereo ({processed_preview.shape}). Cannot apply pan/vol.")
        processed_preview = np.clip(processed_preview, -1.0, 1.0)
        logger.debug(f"Preview generation complete for '{track_name}'. Shape: {processed_preview.shape}")
        return processed_preview
    except Exception as e:
        logger.exception(f"Error generating preview for '{track_name}'")
        st.error(f"Error previewing '{track_name}': {e}")
        return None


# --- REVISED mix_tracks for Sequential Processing (Fades Removed) ---
def mix_tracks(tracks_dict: Dict[TrackID, TrackData], preview: bool = False) -> Tuple[Optional[AudioData], Optional[int]]:
    """Mixes tracks sequentially for memory efficiency."""
    # (Implementation remains the same as v4.4 - Fades already removed)
    logger.info(f"Starting track mixing (Sequential). Preview: {preview}")
    if not tracks_dict:
        logging.warning("Mix called but no tracks.")
        return None, None
    valid_track_ids_for_mix = []
    estimated_processed_lengths = {}
    solo_active = any(t.get("solo", False) for t in tracks_dict.values())
    logger.debug(f"Solo active: {solo_active}")
    logger.info("Step 1: Estimating track lengths after speed changes.")
    for track_id, t_data in tracks_dict.items():
        is_active = t_data.get("solo", False) if solo_active else not t_data.get("mute", False)
        original_audio = t_data.get("original_audio")
        if is_active and original_audio is not None and original_audio.size > 0:
            valid_track_ids_for_mix.append(track_id)
            original_len = len(original_audio)
            speed_factor = t_data.get("speed_factor", 1.0)
            estimated_len = int(original_len / speed_factor) if speed_factor > 0 else original_len
            estimated_processed_lengths[track_id] = estimated_len
            logger.debug(f"Track '{t_data.get('name', track_id)}': Original len={original_len}, Speed={speed_factor}, Estimated len={estimated_len}")
        else:
            logger.debug(f"Skipping track '{t_data.get('name', track_id)}' from length estimation (muted/soloed/no audio).")
    if not valid_track_ids_for_mix:
        logging.warning("No valid tracks found for mixing.")
        return None, None
    target_mix_len_samples = max(estimated_processed_lengths.values()) if estimated_processed_lengths else 0
    logger.info(f"Target mix length based on estimations (pre-looping): {target_mix_len_samples} samples ({target_mix_len_samples / GLOBAL_SR:.2f}s)")
    if preview:
        preview_target_len = int(GLOBAL_SR * MIX_PREVIEW_DURATION_S)
        if target_mix_len_samples > preview_target_len:
            logger.info(f"Preview mode: Limiting mix length from {target_mix_len_samples} to {preview_target_len} samples.")
            target_mix_len_samples = preview_target_len
        process_duration_s = MIX_PREVIEW_DURATION_S + MIX_PREVIEW_PROCESSING_BUFFER_S
        process_samples = int(GLOBAL_SR * process_duration_s)
    if target_mix_len_samples <= 0:
        logging.warning("Mix length <= 0.")
        return None, None
    mix = np.zeros((target_mix_len_samples, 2), dtype=np.float32)
    logger.info(f"Mixing {len(valid_track_ids_for_mix)} tracks sequentially. Initial buffer length: {target_mix_len_samples / GLOBAL_SR:.2f}s")
    actual_max_len_samples = target_mix_len_samples
    processed_segments = {}  # Store processed audio temporarily if needed (only for preview)
    if preview:  # Pre-process preview segments
        for track_id in valid_track_ids_for_mix:
            t_data = tracks_dict[track_id]
            original_audio = t_data.get("original_audio")
            if original_audio is not None and original_audio.size > 0:
                logger.debug(f"Processing PREVIEW segment for track '{t_data.get('name', track_id)}'.")
                segment_samples = min(len(original_audio), process_samples)
                segment = original_audio[:segment_samples].copy()
                processed_segments[track_id] = apply_all_effects(t_data, audio_segment=segment)
            else:
                processed_segments[track_id] = None

    for track_id in valid_track_ids_for_mix:
        t_data = tracks_dict[track_id]
        track_name = t_data.get("name", track_id)
        logger.debug(f"Processing and mixing track: '{track_name}'")
        if preview:
            processed_audio = processed_segments.get(track_id)
        else:
            original_audio = t_data.get("original_audio")
            if original_audio is None or original_audio.size == 0:
                logger.warning(f"Original audio missing for '{track_name}'. Skipping.")
                continue
            processed_audio = apply_all_effects(t_data)  # Process full audio

        if processed_audio is None or processed_audio.size == 0:
            logger.warning(f"Processing failed or empty for '{track_name}'. Skipping.")
            continue
        actual_processed_len = len(processed_audio)
        logger.debug(f"Track '{track_name}': Actual processed length = {actual_processed_len}")
        final_audio_for_track = processed_audio
        should_loop = t_data.get("loop_to_fit", False)
        if not preview and should_loop:
            if target_mix_len_samples > 0 and actual_processed_len > 0 and actual_processed_len < target_mix_len_samples:
                logger.info(f"Looping track '{track_name}' from {actual_processed_len} to {target_mix_len_samples} samples.")
                n_repeats = target_mix_len_samples // actual_processed_len
                remainder = target_mix_len_samples % actual_processed_len
                looped_list = [processed_audio] * n_repeats
                if remainder > 0:
                    looped_list.append(processed_audio[:remainder])
                try:
                    final_audio_for_track = np.concatenate(looped_list, axis=0)
                    actual_max_len_samples = max(actual_max_len_samples, len(final_audio_for_track))
                    logger.debug(f"Looping complete for '{track_name}'. New length: {len(final_audio_for_track)}")
                except ValueError as e_concat:
                    logger.error(f"Error concatenating looped audio for '{track_name}': {e_concat}. Skipping loop.")
                    final_audio_for_track = processed_audio
            else:
                logger.debug(f"Looping not needed for '{track_name}' (curr: {actual_processed_len}, target: {target_mix_len_samples})")
        if len(mix) < actual_max_len_samples:
            logger.warning(f"Resizing mix buffer from {len(mix)} to {actual_max_len_samples} due to looping.")
            mix = np.pad(mix, ((0, actual_max_len_samples - len(mix)), (0, 0)), mode="constant")
            target_mix_len_samples = actual_max_len_samples
        current_len = len(final_audio_for_track)
        if current_len < target_mix_len_samples:
            audio_adjusted = np.pad(final_audio_for_track, ((0, target_mix_len_samples - current_len), (0, 0)), mode="constant")
        elif current_len > target_mix_len_samples:
            audio_adjusted = final_audio_for_track[:target_mix_len_samples, :]
        else:
            audio_adjusted = final_audio_for_track
        pan = t_data.get("pan", 0.0)
        vol = t_data.get("volume", 1.0)
        logger.debug(f"Track '{track_name}': Applying vol={vol:.2f}, pan={pan:.2f}")
        pan_rad = (pan + 1) * np.pi / 4
        left_gain = vol * np.cos(pan_rad)
        right_gain = vol * np.sin(pan_rad)
        if audio_adjusted.ndim == 2 and audio_adjusted.shape[1] == 2:
            mix[:, 0] += audio_adjusted[:, 0] * left_gain
            mix[:, 1] += audio_adjusted[:, 1] * right_gain
        elif audio_adjusted.ndim == 1:
            logger.warning(f"Track '{track_name}' is mono during mixing.")
            mix[:, 0] += audio_adjusted * vol * 0.707
            mix[:, 1] += audio_adjusted * vol * 0.707
        logger.debug(f"Added track '{track_name}' to mix buffer.")
        del processed_audio, final_audio_for_track, audio_adjusted  # Clear memory
    final_mix = np.clip(mix, -1.0, 1.0)
    # --- REMOVED Master Fades ---
    logger.info("Mixing complete.")
    return final_mix.astype(np.float32), len(final_mix)  # Return audio and length


# --- Helper to read text files ---
# (read_text_file function remains the same as v2.12 with enhanced checks)
def read_text_file(uploaded_file: UploadedFile) -> Optional[str]:
    """Reads content from uploaded txt or docx file with enhanced error handling."""
    if uploaded_file is None:
        logger.error("read_text_file called with None object.")
        return None
    file_name = uploaded_file.name
    logger.info(f"Attempting to read file: {file_name}, Type: {uploaded_file.type}")
    try:
        if file_name.lower().endswith(".txt"):
            logger.debug(f"Reading .txt file: {file_name}")
            try:
                return uploaded_file.read().decode("utf-8", errors="replace")
            except Exception as e_decode:
                logger.exception(f"Error decoding .txt file {file_name}")
                st.error(f"Error decoding: {e_decode}")
                return None
        elif file_name.lower().endswith(".docx"):
            logger.debug(f"Processing .docx file: {file_name}")
            try:
                uploaded_file.seek(0)
                document = docx.Document(uploaded_file)
                logger.debug(f"docx created. Found {len(document.paragraphs)} paragraphs.")
                para_texts = []
                for i, para in enumerate(document.paragraphs):
                    if para is None:
                        logger.warning(f"Para {i} in {file_name} is None.")
                        continue
                    para_text = getattr(para, "text", "")  # Safely get text
                    if para_text:
                        para_texts.append(para_text)
                logger.debug(f"Extracted {len(para_texts)} non-empty paragraphs.")
                return "\n".join(para_texts)
            except Exception as e_docx:
                logger.exception(f"Error processing docx: {file_name}")
                st.error(f"Error reading Word doc: {e_docx}")
                return None
        else:
            st.error(f"Unsupported file type: {uploaded_file.type}.")
            logger.error(f"Unsupported type: {uploaded_file.type}")
            return None
    except Exception as e_outer:
        logger.exception(f"Outer error reading file: {uploaded_file.name}")
        st.error(f"Error reading file: {e_outer}")
        return None


# ==========================================
# 2. Application State Management - PREVIEW CACHING
# ==========================================
# (AppState class remains the same as v4.6)
class AppState:
    """Manages the application state using non-destructive approach + preview path + track type + preview hash."""

    STATE_KEY = "tracks_non_destructive_v4"  # Use a new key

    def __init__(self):
        if self.STATE_KEY not in st.session_state:
            logger.info(f"Initializing '{self.STATE_KEY}'.")
            st.session_state[self.STATE_KEY] = {}
        default_keys = AppState.get_default_track_params().keys()
        tracks_dict = self.get_all_tracks()
        for track_id, track_data in list(tracks_dict.items()):
            changed = False
            if "processed_audio" in track_data:
                del st.session_state[self.STATE_KEY][track_id]["processed_audio"]
                logger.debug(f"Removed old 'processed_audio' key for {track_id}")
                changed = True
            for key, default_value in AppState.get_default_track_params().items():
                if key not in track_data:
                    logger.warning(f"Track {track_id} missing key '{key}', adding default.")
                    st.session_state[self.STATE_KEY][track_id][key] = default_value
                    changed = True
            if track_data.get("preview_settings_hash") is None and track_data.get("preview_temp_file_path"):
                old_path = track_data["preview_temp_file_path"]
                if old_path and os.path.exists(old_path):
                    try:
                        os.remove(old_path)
                        logger.info(f"Cleaned up old preview file {old_path} due to missing hash.")
                    except OSError as e:
                        logger.warning(f"Could not clean up old preview file {old_path}: {e}")
                st.session_state[self.STATE_KEY][track_id]["preview_temp_file_path"] = None

    @staticmethod
    def get_default_track_params() -> TrackData:
        """Returns default parameters including preview hash."""
        return {
            "original_audio": np.zeros((0, 2), dtype=np.float32),
            "sr": GLOBAL_SR,
            "name": "New Track",
            "track_type": TRACK_TYPE_OTHER,
            "volume": 1.0,
            "mute": False,
            "solo": False,
            "speed_factor": 1.0,
            "pitch_shift": 0,
            "pan": 0.0,
            "filter_type": "off",
            "filter_cutoff": 8000.0,
            "loop_to_fit": False,
            "reverse_audio": False,
            "preview_temp_file_path": None,  # Path to the current preview file
            "preview_settings_hash": None,  # Hash of settings used for current preview
            "update_counter": 0,
        }

    def _get_tracks_dict(self) -> Dict[TrackID, TrackData]:
        return st.session_state.get(self.STATE_KEY, {})

    def get_all_tracks(self) -> Dict[TrackID, TrackData]:
        return self._get_tracks_dict()

    def get_track(self, track_id: TrackID) -> Optional[TrackData]:
        return self._get_tracks_dict().get(track_id)

    def add_track(self, track_id: TrackID, track_data: TrackData, track_type: TrackType = TRACK_TYPE_OTHER):
        """Adds a new track, storing original audio, parameters, and type."""
        if not isinstance(track_data, dict):
            logger.error(f"Invalid track data type {track_id}")
            return
        default_params = AppState.get_default_track_params()
        final_track_data = {key: track_data.get(key, default_value) for key, default_value in default_params.items()}
        if "original_audio" not in track_data or track_data["original_audio"] is None:
            logger.error(f"Attempted add track {track_id} without original_audio.")
            return
        final_track_data["original_audio"] = track_data["original_audio"]
        final_track_data["sr"] = track_data.get("sr", GLOBAL_SR)
        final_track_data["name"] = track_data.get("name", default_params["name"])
        final_track_data["track_type"] = track_type  # Set the type
        final_track_data["preview_temp_file_path"] = None  # Initialize preview path
        final_track_data["preview_settings_hash"] = None  # Initialize hash
        st.session_state[self.STATE_KEY][track_id] = final_track_data
        logger.info(f"Added track ID: {track_id}, Name: '{final_track_data.get('name', 'N/A')}', Type: {track_type}")

    def delete_track(self, track_id: TrackID):
        """Deletes a track and its associated preview temp file."""
        tracks = self._get_tracks_dict()
        if track_id in tracks:
            track_name = tracks[track_id].get("name", "N/A")
            preview_path = tracks[track_id].get("preview_temp_file_path")
            if preview_path and os.path.exists(preview_path):
                try:
                    os.remove(preview_path)
                    logger.info(f"Deleted preview temp file '{preview_path}' for track {track_id}")
                except OSError as e:
                    logger.warning(f"Failed delete preview temp file '{preview_path}': {e}")
            del st.session_state[self.STATE_KEY][track_id]
            logger.info(f"Deleted track ID: {track_id}, Name: '{track_name}'")
            return True
        else:
            logger.warning(f"Attempted delete non-existent track {track_id}")
            return False

    def update_track_param(self, track_id: TrackID, param_name: str, value: Any):
        """Updates a specific parameter for a given track."""
        tracks = self._get_tracks_dict()
        if track_id in tracks:
            # Invalidate hash if a parameter affecting preview changes
            preview_params = ["volume", "speed_factor", "pitch_shift", "pan", "filter_type", "filter_cutoff", "reverse_audio"]
            if param_name in preview_params and tracks[track_id].get(param_name) != value:
                logger.debug(f"Parameter '{param_name}' changed for track {track_id}, invalidating preview hash.")
                st.session_state[self.STATE_KEY][track_id]["preview_settings_hash"] = None  # Invalidate hash

            if param_name in AppState.get_default_track_params() and param_name != "original_audio":
                st.session_state[self.STATE_KEY][track_id][param_name] = value
            elif param_name == "original_audio":
                logger.error(f"Attempted update original_audio directly for {track_id}.")
            else:
                logger.warning(f"Attempted update invalid param '{param_name}' for {track_id}")
        else:
            logger.warning(f"Attempted update param for non-existent track {track_id}")

    def increment_update_counter(self, track_id: TrackID):
        """Increments the update counter for a track."""
        tracks = self._get_tracks_dict()
        if track_id in tracks:
            current = tracks[track_id].get("update_counter", 0)
            new = current + 1
            st.session_state[self.STATE_KEY][track_id]["update_counter"] = new
            logger.debug(f"Incr update_counter for {track_id} to {new}")
        else:
            logger.warning(f"Attempted incr counter for non-existent track {track_id}")

    def get_loaded_track_names(self) -> List[str]:
        return [t.get("name") for t in self.get_all_tracks().values() if t.get("name")]


# ==========================================
# 3. TTS Generation (Class Wrapper) - WITH CHUNKING
# ==========================================
# (TTSGenerator class remains the same as v2.15 - Fixed Spinner Error)
class TTSGenerator:
    """Handles Text-to-Speech generation using pyttsx3 with chunking."""

    def __init__(self, chunk_size: int = TTS_CHUNK_SIZE):
        self.engine = None
        self.rate = 200
        self.volume = 1.0
        self.chunk_size = chunk_size
        logger.debug(f"TTSGenerator initialized chunk={chunk_size}, rate={self.rate}")

    def _init_engine(self):
        try:
            logger.debug("Initializing pyttsx3 engine for generation.")
            self.engine = pyttsx3.init()
            if self.engine is None:
                raise RuntimeError("pyttsx3.init() returned None")
            self.engine.setProperty("rate", self.rate)
            self.engine.setProperty("volume", self.volume)
        except Exception as e:
            logger.exception("Failed init pyttsx3.")
            self.engine = None
            raise RuntimeError("TTS Engine init failed") from e

    def generate(self, text: str) -> Tuple[Optional[AudioData], Optional[SampleRate]]:
        logger.info(f"Starting TTS generation for text length: {len(text)} chars.")
        if not text:
            logger.warning("Empty TTS text provided.")
            return None, None
        temp_chunk_files = []
        audio_chunks = []
        final_sr = None
        progress_placeholder = st.empty()
        try:
            self._init_engine()
            if self.engine is None:
                st.error("TTS Engine could not be initialized.")
                return None, None
            logger.debug(f"Wrapping text into chunks of size: {self.chunk_size}")
            chunks = textwrap.wrap(text, self.chunk_size, break_long_words=True, replace_whitespace=False, drop_whitespace=True)
            num_chunks = len(chunks)
            logger.info(f"Split text into {num_chunks} chunks.")
            with st.spinner(f"Synthesizing {num_chunks} audio chunks..."):
                for i, chunk in enumerate(chunks):
                    if not chunk.strip():
                        logger.debug(f"Skipping empty chunk {i + 1}/{num_chunks}")
                        continue
                    chunk_start_time = time.time()
                    progress_placeholder.text(f"Synthesizing audio chunk {i + 1}/{num_chunks}...")
                    temp_chunk_path = None
                    try:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_chunk{i}.wav", mode="wb") as tmp:
                            temp_chunk_path = tmp.name
                        temp_chunk_files.append(temp_chunk_path)
                        logger.debug(f"Synthesizing chunk {i + 1}/{num_chunks} to {temp_chunk_path}...")
                        self.engine.save_to_file(chunk, temp_chunk_path)
                        self.engine.runAndWait()
                        chunk_end_time = time.time()
                        logger.debug(f"Chunk {i + 1} synthesized in {chunk_end_time - chunk_start_time:.2f}s.")
                        if not os.path.exists(temp_chunk_path) or os.path.getsize(temp_chunk_path) == 0:
                            logger.error(f"TTS engine failed create non-empty file chunk {i + 1}: {temp_chunk_path}")
                            st.error(f"TTS engine failed for chunk {i + 1}.")
                            if temp_chunk_path in temp_chunk_files:
                                temp_chunk_files.remove(temp_chunk_path)
                            temp_chunk_path = None
                            continue
                    except Exception as e_chunk:
                        logger.exception(f"Failed synthesize chunk {i + 1}")
                        st.error(f"Error processing chunk {i + 1}: {e_chunk}")
                        if temp_chunk_path in temp_chunk_files:
                            try:
                                temp_chunk_files.remove(temp_chunk_path)
                                logger.debug(f"Removed problematic chunk path {temp_chunk_path}.")
                            except ValueError:
                                logger.warning(f"Could not remove {temp_chunk_path} from list.")
                        continue
                progress_placeholder.text("Combining audio chunks...")
                logger.info("Synthesized all chunks. Concatenating.")
                for i, file_path in enumerate(temp_chunk_files):
                    if not file_path or not os.path.exists(file_path):
                        logger.warning(f"Temp chunk file {file_path} missing, skipping.")
                        continue
                    try:
                        logger.debug(f"Loading chunk file: {file_path}")
                        audio_data, sr = sf.read(file_path, dtype="float32", always_2d=True)
                        if audio_data.shape[1] == 1:
                            audio_data = np.concatenate([audio_data, audio_data], axis=1)
                        if final_sr is None:
                            final_sr = sr
                            logger.debug(f"Detected SR: {final_sr} Hz")
                        elif sr != final_sr:
                            logger.warning(f"SR mismatch! Expected {final_sr}, got {sr}. Resampling chunk.")
                            if audio_data.size > 0:
                                audio_data = librosa.resample(audio_data.T, orig_sr=sr, target_sr=final_sr).T.astype(np.float32)
                            else:
                                continue
                        audio_chunks.append(audio_data)
                        logger.debug(f"Loaded chunk {i + 1} shape {audio_data.shape}")
                    except Exception as e_load:
                        logger.exception(f"Failed load chunk {file_path}")
                        st.error(f"Error loading chunk {i + 1}: {e_load}")
            progress_placeholder.empty()
            if not audio_chunks:
                logger.error("No valid chunks loaded.")
                st.error("Failed process chunks.")
                return None, None
            logger.info(f"Concatenating {len(audio_chunks)} audio chunks...")
            final_audio = np.concatenate(audio_chunks, axis=0)
            if final_sr is None:
                logger.error("Could not determine SR.")
                st.error("Failed determine SR.")
                return None, None
            if final_sr != GLOBAL_SR:
                logger.info(f"Resampling final audio from {final_sr} Hz to {GLOBAL_SR} Hz.")
                if final_audio.size > 0:
                    final_audio = librosa.resample(final_audio.T, orig_sr=final_sr, target_sr=GLOBAL_SR).T.astype(np.float32)
                final_sr = GLOBAL_SR
            logger.info(f"TTS generation complete. Final shape: {final_audio.shape}, SR: {final_sr}")
            return final_audio, final_sr
        except Exception as e:
            progress_placeholder.empty()
            logger.exception("TTS Gen Failed.")
            st.error(f"TTS Gen Failed: {e}")
            return None, None
        finally:
            progress_placeholder.empty()
            logger.debug(f"Cleaning up {len(temp_chunk_files)} temp chunk files...")
            cleaned_count = 0
            for file_path in temp_chunk_files:
                if file_path and os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                        cleaned_count += 1
                    except OSError as e_os:
                        logger.warning(f"Could not delete temp TTS chunk {file_path}: {e_os}")
            logger.debug(f"Cleaned up {cleaned_count} temp files.")


# ==========================================
# 4. UI Management (Class) - PREVIEW CACHING
# ==========================================
# (UIManager class remains the same as v4.6)
class UIManager:
    """Handles rendering Streamlit UI components using button-driven previews with caching."""

    def __init__(self, app_state: AppState, tts_generator: TTSGenerator):
        self.app_state = app_state
        self.tts_generator = tts_generator
        logger.debug("UIManager initialized.")

    # --- Helper function to calculate preview settings hash ---
    def _calculate_preview_hash(self, track_data: TrackData) -> int:
        """Calculates a hash based on settings relevant to the preview."""
        params_to_hash = (
            track_data.get("volume", 1.0),
            track_data.get("speed_factor", 1.0),
            track_data.get("pitch_shift", 0),
            track_data.get("pan", 0.0),
            track_data.get("filter_type", "off"),
            track_data.get("filter_cutoff", 8000.0),
            track_data.get("reverse_audio", False),
            track_data.get("original_audio").size if track_data.get("original_audio") is not None else 0,
        )
        settings_hash = hash(params_to_hash)
        logger.debug(f"Calculated preview hash for track '{track_data.get('name', 'N/A')}': {settings_hash}")
        return settings_hash

    # (render_sidebar and its helpers remain the same as v4.0)
    def render_sidebar(self):
        with st.sidebar:
            logo_path = os.path.join("assets", "logo.png")
            if os.path.exists(logo_path):
                st.image(logo_path, width=200)
            else:
                st.header("MindMorph")
            st.caption("Subliminal Audio Editor")
            st.markdown("---")
            st.markdown("### STEP 1: Add Audio Layers")
            st.markdown("Use options below to add sounds.")
            self._render_uploader()
            st.divider()
            self._render_affirmation_inputs()
            st.divider()  # Uses new method
            self._render_frequency_generators()
            st.divider()  # Combined Freq Gens
            self._render_background_generators()
            st.markdown("---")  # Background Noise
            # --- Project Save/Load ---
            st.subheader("💾 Project")
            self._render_save_load()
            st.markdown("---")
            # -------------------------
            st.info("Edit track details in the main panel.")

    def _render_uploader(self):
        st.subheader("📁 Upload Audio File(s)")
        st.caption("Upload music, voice recordings, or other sounds.")
        uploaded_files = st.file_uploader(
            "Select audio files",
            type=["wav", "mp3", "ogg", "flac"],
            accept_multiple_files=True,
            key="upload_files_key",
            label_visibility="collapsed",
            help="Select one or more audio files.",
        )
        loaded_track_names = self.app_state.get_loaded_track_names()
        if uploaded_files:
            for file in uploaded_files:
                if file.name not in loaded_track_names:
                    logger.info(f"Processing upload: {file.name}")
                    with st.spinner(f"Loading {file.name}..."):
                        audio, sr = load_audio(file, target_sr=GLOBAL_SR)
                    if audio.size > 0:
                        track_id = str(uuid.uuid4())
                        track_params = AppState.get_default_track_params()
                        # Deduce likely type based on name? Simple heuristic for now.
                        track_type = TRACK_TYPE_VOICE if "voice" in file.name.lower() or "record" in file.name.lower() else TRACK_TYPE_BACKGROUND
                        track_params.update(
                            {
                                "original_audio": audio,
                                "sr": sr,
                                "name": file.name,
                                "source_type": "upload",  # Store source type
                                "original_filename": file.name,  # Store original filename
                            }
                        )
                        self.app_state.add_track(track_id, track_params, track_type=track_type)
                        st.success(f"Loaded '{file.name}'")
                        loaded_track_names.append(file.name)
                    else:
                        logger.warning(f"Skipped empty upload: {file.name}")
                        st.warning(f"Skipped empty: {file.name}")

    def _render_affirmation_inputs(self):
        """Renders options for adding affirmations (TTS, File Upload, Record Placeholder)."""
        st.subheader("🗣️ Add Affirmations")
        st.caption("Uses system default TTS voice. Voice selection depends on OS/browser.")
        tab1, tab2, tab3 = st.tabs(["Type Text", "Upload File", "Record Audio"])
        with tab1:
            st.caption("Type or paste affirmations below.")
            affirmation_text = st.text_area(
                "Affirmations (one per line)", height=150, key="affirmation_text_area", label_visibility="collapsed", help="Enter each affirmation on a new line."
            )
            if st.button(
                "Generate Affirmation Track", key="generate_tts_text_key", use_container_width=True, type="primary", help="Convert the text above to a spoken audio track."
            ):
                if affirmation_text:
                    self._generate_tts_track(affirmation_text, "Affirmations (Text)")
                else:
                    st.warning("Please enter text in the text area first.")
        with tab2:
            st.caption("Upload a .txt or .docx file containing affirmations.")
            uploaded_file = st.file_uploader(
                "Upload Affirmation File", type=["txt", "docx"], key="affirmation_file_uploader", label_visibility="collapsed", help="Select text/Word document."
            )
            if st.button(
                "Generate Track from File",
                key="generate_tts_file_key",
                use_container_width=True,
                type="primary",
                help="Read the uploaded file and convert its text content to a spoken audio track.",
            ):
                if uploaded_file:
                    with st.spinner(f"Reading {uploaded_file.name}..."):
                        text = read_text_file(uploaded_file)
                    if text is not None:
                        if text.strip():
                            self._generate_tts_track(text, f"Affirmations ({uploaded_file.name})")
                        else:
                            st.warning(f"File '{uploaded_file.name}' has no readable text.")
                            logger.warning(f"File '{uploaded_file.name}' read empty.")
                    else:
                        logger.error("Failed read uploaded file for TTS.")
                else:
                    st.warning("Please upload file.")
        with tab3:
            st.caption("Record your own voice for affirmations.")
            st.info("🎙️ Audio recording feature coming soon!")
            st.markdown("Use 'Upload Audio File(s)' for now.")
            st.button("Start Recording", key="start_record_key", disabled=True, use_container_width=True)

    def _generate_tts_track(self, text_content: str, track_name: str):
        """Helper to generate TTS and add track, tagged as Affirmation."""
        audio, sr = self.tts_generator.generate(text_content)  # Spinner inside generate
        if audio is not None and sr is not None:
            track_id = str(uuid.uuid4())
            track_params = AppState.get_default_track_params()
            track_params.update(
                {
                    "original_audio": audio,
                    "sr": sr,
                    "name": track_name,
                    "source_type": "tts",  # Store source type
                    "tts_text": text_content,  # Store the text used
                }
            )
            self.app_state.add_track(track_id, track_params, track_type=TRACK_TYPE_AFFIRMATION)  # Tag as affirmation
            st.success(f"'{track_name}' track generated!")
            st.toast("Affirmation track added!", icon="✅")

    def _render_frequency_generators(self):
        """Renders options for generating Binaural Beats and Solfeggio Tones."""
        st.subheader("🧠✨ Add Frequencies / Tones")
        gen_type = st.radio(
            "Select Type:", ["Binaural Beats", "Solfeggio Tones", "Isochronic Tones", "Presets"], key="freq_gen_type", horizontal=True, label_visibility="collapsed"
        )  # Added Isochronic

        if gen_type == "Binaural Beats":
            st.markdown("<small>Generates stereo tones potentially inducing brainwave states (requires headphones).</small>", unsafe_allow_html=True)
            bb_cols = st.columns(2)
            bb_duration = bb_cols[0].number_input("Duration (s)", 1, 3600, 60, 1, key="bb_duration", help="Length in seconds.")
            bb_vol = bb_cols[1].slider("Volume##BB", 0.0, 1.0, 0.3, 0.05, key="bb_volume", help="Loudness (0.0 to 1.0). Usually kept low.")
            bb_fcols = st.columns(2)
            bb_fleft = bb_fcols[0].number_input("Left Freq (Hz)", 20, 1000, 200, 1, key="bb_freq_left", help="Left ear frequency.")
            bb_fright = bb_fcols[1].number_input("Right Freq (Hz)", 20, 1000, 210, 1, key="bb_freq_right", help="Right ear frequency.")
            if st.button("Generate Binaural Track", key="generate_bb", help="Create track with these settings."):
                with st.spinner("Generating..."):
                    audio = generate_binaural_beats(bb_duration, bb_fleft, bb_fright, GLOBAL_SR, bb_vol)
                track_id = str(uuid.uuid4())
                track_params = AppState.get_default_track_params()
                track_params.update(
                    {
                        "original_audio": audio,
                        "sr": GLOBAL_SR,
                        "name": f"Binaural {bb_fleft}/{bb_fright}Hz",
                        "source_type": "binaural",
                        "gen_duration": bb_duration,
                        "gen_freq_left": bb_fleft,
                        "gen_freq_right": bb_fright,
                        "gen_volume": bb_vol,  # Store generation params
                    }
                )
                self.app_state.add_track(track_id, track_params, track_type=TRACK_TYPE_FREQUENCY)
                st.success("Binaural Beats generated!")

        elif gen_type == "Solfeggio Tones":
            st.markdown("<small>Generates pure tones based on historical Solfeggio frequencies.</small>", unsafe_allow_html=True)
            freqs = [174, 285, 396, 417, 528, 639, 741, 852, 963]
            cols = st.columns(2)
            freq = cols[0].selectbox("Frequency (Hz)", freqs, index=4, key="solf_freq", help="Select Solfeggio frequency.")
            duration = cols[1].number_input("Duration (s)##Solf", 1, 3600, 60, 1, key="solf_duration", help="Length in seconds.")
            vol = st.slider("Volume##Solf", 0.0, 1.0, 0.3, 0.05, key="solf_volume", help="Loudness (0.0 to 1.0). Usually kept low.")
            if st.button("Generate Solfeggio Track", key="generate_solf", help="Create track with this tone."):
                with st.spinner("Generating..."):
                    audio = generate_solfeggio_frequency(duration, freq, GLOBAL_SR, vol)
                track_id = str(uuid.uuid4())
                track_params = AppState.get_default_track_params()
                track_params.update(
                    {
                        "original_audio": audio,
                        "sr": GLOBAL_SR,
                        "name": f"Solfeggio {freq}Hz",
                        "source_type": "solfeggio",
                        "gen_duration": duration,
                        "gen_freq": freq,
                        "gen_volume": vol,
                    }
                )
                self.app_state.add_track(track_id, track_params, track_type=TRACK_TYPE_FREQUENCY)
                st.success(f"Solfeggio {freq}Hz generated!")

        elif gen_type == "Isochronic Tones":
            st.markdown("<small>Generates rhythmic pulses of a single tone (headphones not required).</small>", unsafe_allow_html=True)
            iso_cols = st.columns(2)
            iso_duration = iso_cols[0].number_input("Duration (s)##Iso", 1, 3600, 60, 1, key="iso_duration", help="Length in seconds.")
            iso_vol = iso_cols[1].slider("Volume##Iso", 0.0, 1.0, 0.4, 0.05, key="iso_volume", help="Loudness (0.0 to 1.0).")
            iso_fcols = st.columns(2)
            iso_carrier = iso_fcols[0].number_input("Carrier Freq (Hz)", 20, 1000, 150, 1, key="iso_carrier", help="The base tone frequency.")
            iso_pulse = iso_fcols[1].number_input("Pulse Freq (Hz)", 1, 40, 10, 1, key="iso_pulse", help="How many times per second the tone pulses.")
            if st.button("Generate Isochronic Track", key="generate_iso", help="Create track with these settings."):
                with st.spinner("Generating..."):
                    audio = generate_isochronic_tones(iso_duration, iso_carrier, iso_pulse, GLOBAL_SR, iso_vol)
                track_id = str(uuid.uuid4())
                track_params = AppState.get_default_track_params()
                track_params.update(
                    {
                        "original_audio": audio,
                        "sr": GLOBAL_SR,
                        "name": f"Isochronic {iso_carrier}/{iso_pulse}Hz",
                        "source_type": "isochronic",
                        "gen_duration": iso_duration,
                        "gen_carrier_freq": iso_carrier,
                        "gen_pulse_freq": iso_pulse,
                        "gen_volume": iso_vol,
                    }
                )
                self.app_state.add_track(track_id, track_params, track_type=TRACK_TYPE_FREQUENCY)
                st.success("Isochronic Tones generated!")

        elif gen_type == "Presets":
            st.markdown("<small>Generate common frequency patterns.</small>", unsafe_allow_html=True)
            preset_options = {
                "Focus (Alpha 10Hz)": {"type": "binaural", "f_left": 200, "f_right": 210},
                "Relaxation (Theta 5Hz)": {"type": "binaural", "f_left": 150, "f_right": 155},
                "Deep Sleep (Delta 2Hz)": {"type": "binaural", "f_left": 100, "f_right": 102},
                "Love Freq (Solfeggio 528Hz)": {"type": "solfeggio", "freq": 528},
                "Miracle Tone (Solfeggio 417Hz)": {"type": "solfeggio", "freq": 417},
            }
            preset_name = st.selectbox("Select Preset:", list(preset_options.keys()), key="freq_preset_select")
            preset_data = preset_options[preset_name]
            cols_preset = st.columns(2)
            preset_duration = cols_preset[0].number_input("Duration (s)##Preset", 1, 3600, 60, 1, key="preset_duration")
            preset_vol = cols_preset[1].slider("Volume##Preset", 0.0, 1.0, 0.2, 0.05, key="preset_volume", help="Volume (usually kept low)")

            if st.button(f"Generate '{preset_name}' Track", key="generate_preset_freq"):
                with st.spinner("Generating preset frequency..."):
                    audio = None
                    gen_params = {}
                    if preset_data["type"] == "binaural":
                        audio = generate_binaural_beats(preset_duration, preset_data["f_left"], preset_data["f_right"], GLOBAL_SR, preset_vol)
                        gen_params = {
                            "source_type": "binaural",
                            "gen_duration": preset_duration,
                            "gen_freq_left": preset_data["f_left"],
                            "gen_freq_right": preset_data["f_right"],
                            "gen_volume": preset_vol,
                        }
                    elif preset_data["type"] == "solfeggio":
                        audio = generate_solfeggio_frequency(preset_duration, preset_data["freq"], GLOBAL_SR, preset_vol)
                        gen_params = {"source_type": "solfeggio", "gen_duration": preset_duration, "gen_freq": preset_data["freq"], "gen_volume": preset_vol}

                    if audio is not None:
                        track_id = str(uuid.uuid4())
                        track_params = AppState.get_default_track_params()
                        track_params.update({"original_audio": audio, "sr": GLOBAL_SR, "name": preset_name})
                        track_params.update(gen_params)  # Add generation parameters
                        self.app_state.add_track(track_id, track_params, track_type=TRACK_TYPE_FREQUENCY)
                        st.success(f"'{preset_name}' track generated!")
                    else:
                        st.error("Failed to generate audio for selected preset.")

    def _render_background_generators(self):
        st.subheader("🎵 Add Background Noise")
        noise_options = ["White Noise", "Pink Noise", "Brown Noise"]
        noise_type = st.selectbox("Select Noise Type:", noise_options, key="noise_type_select")
        cols_noise = st.columns(2)
        noise_duration = cols_noise[0].number_input("Duration (s)##Noise", 10, 7200, 300, 10, key="noise_duration", help="Length in seconds (will loop if shorter than project).")
        noise_vol = cols_noise[1].slider("Volume##Noise", 0.0, 1.0, 0.5, 0.05, key="noise_volume", help="Loudness (0.0 to 1.0).")

        if st.button(f"Generate {noise_type} Track", key="generate_noise"):
            with st.spinner(f"Generating {noise_type}..."):
                audio = generate_noise(noise_type, noise_duration, GLOBAL_SR, noise_vol)
                if audio is not None:
                    track_id = str(uuid.uuid4())
                    track_params = AppState.get_default_track_params()
                    track_params.update(
                        {
                            "original_audio": audio,
                            "sr": GLOBAL_SR,
                            "name": noise_type,
                            "loop_to_fit": True,
                            "source_type": "noise",
                            "gen_noise_type": noise_type,
                            "gen_duration": noise_duration,
                            "gen_volume": noise_vol,  # Store generation params
                        }
                    )
                    self.app_state.add_track(track_id, track_params, track_type=TRACK_TYPE_BACKGROUND)
                    st.success(f"{noise_type} track generated!")
                else:
                    st.error(f"Failed to generate {noise_type}.")
        st.caption("More noise types (Rain etc.) coming soon!")

    # --- NEW: Render Save/Load Buttons ---
    def _render_save_load(self):
        """Renders project save and load components."""
        st.markdown("**Save/Load Project**")
        st.caption("Save project configuration (track settings, sources) to a `.mindmorph` file. Audio data itself is not saved.")

        # --- Save Button ---
        project_data_str = ""
        tracks = self.app_state.get_all_tracks()
        if tracks:
            try:
                serializable_tracks = {}
                for track_id, track_data in tracks.items():
                    save_data = track_data.copy()
                    # Remove non-serializable audio data, keep metadata
                    if "original_audio" in save_data:
                        del save_data["original_audio"]
                    if "preview_temp_file_path" in save_data:
                        del save_data["preview_temp_file_path"]  # Don't save temp path
                    if "preview_settings_hash" in save_data:
                        del save_data["preview_settings_hash"]  # Don't save hash
                    # Ensure necessary source info is present
                    if "source_type" not in save_data:
                        save_data["source_type"] = "unknown"
                    serializable_tracks[track_id] = save_data

                project_file_content = {"version": PROJECT_FILE_VERSION, "tracks": serializable_tracks}
                project_data_str = json.dumps(project_file_content, indent=2)
            except Exception as e:
                logger.error(f"Error preparing project data for saving: {e}")
                st.error("Could not prepare project data for saving.")

        save_disabled = not bool(tracks) or not project_data_str
        st.download_button(
            label="💾 Save Project File",
            data=project_data_str,
            file_name="my_subliminal_project.mindmorph",
            mime="application/json",
            key="save_project_button",
            help="Saves the current track list and settings.",
            use_container_width=True,
            disabled=save_disabled,
        )

        # --- Load Button ---
        uploaded_project_file = st.file_uploader(
            "⬆️ Load Project File (.mindmorph)",
            type=["mindmorph", "json"],
            key="load_project_uploader",
            accept_multiple_files=False,
            help="Load a previously saved project configuration. This will replace the current project.",
        )

        if uploaded_project_file is not None:
            # Store uploaded file in session state to handle after rerun
            st.session_state.uploaded_project_file_data = uploaded_project_file.getvalue()
            st.session_state.project_load_requested = True
            # Clear the uploader state by setting the key to random after processing
            st.session_state.load_project_uploader = str(uuid.uuid4())
            logger.info(f"Project file uploaded: {uploaded_project_file.name}. Requesting load.")
            st.rerun()  # Rerun to trigger the loading logic in main()

    # --- (render_tracks_editor, _render_track_main_col, _render_track_controls_col remain the same as v4.6) ---
    def render_tracks_editor(self):
        """Renders the main editor area with all tracks using previews."""
        st.header("🎚️ Tracks Editor")
        tracks = self.app_state.get_all_tracks()
        if not tracks:
            if "welcome_message_shown" in st.session_state:
                with st.container(border=True):
                    st.markdown("#### ✨ Your Project is Empty!")
                    st.markdown("Ready to start creating? Use the **sidebar on the left** (👈) to:")
                    st.markdown("- Add your **🗣️ Affirmations** (Type/Upload/Record)")
                    st.markdown("- Add a **🎵 Background** sound (Upload/Generate)")
                    st.markdown("- Optionally add **🧠✨ Tones** (Binaural/Solfeggio/Presets)")
                    st.markdown("---")
                    st.markdown("Once you add a track, the editor controls will appear here.")
            return
        st.markdown("Adjust settings below. Click **'Update Preview'** to refresh the 60s preview with current settings.")
        st.divider()
        track_ids_to_delete = []
        logger.debug(f"Rendering editor for {len(tracks)} tracks.")
        for track_id, track_data in list(tracks.items()):
            if track_id not in self.app_state.get_all_tracks():
                continue
            track_name = track_data.get("name", "Unnamed")
            track_type_icon = track_data.get("track_type", TRACK_TYPE_OTHER).split(" ")[0]  # Get icon
            # --- Display message if uploaded source is missing ---
            expander_label = f"{track_type_icon} Track: **{track_name}** (`{track_id[:6]}`)"
            if track_data.get("source_type") == "upload" and track_data.get("original_audio") is None:
                expander_label += " ⚠️ Missing Source File"
            # ---------------------------------------------------
            with st.expander(expander_label, expanded=True):
                logger.debug(f"Rendering expander for: '{track_name}' ({track_id}) Type: {track_data.get('track_type')}")
                col_main, col_controls = st.columns([3, 1])
                self._render_track_main_col(track_id, track_data, col_main)
                deleted = self._render_track_controls_col(track_id, track_data, col_controls)
                if deleted:
                    track_ids_to_delete.append(track_id)
        if track_ids_to_delete:
            deleted_count = 0
            for tid in track_ids_to_delete:
                if self.app_state.delete_track(tid):
                    deleted_count += 1
            if deleted_count > 0:
                st.toast(f"Deleted {deleted_count} track(s).")
                st.rerun()

    def _render_track_main_col(self, track_id: TrackID, track_data: TrackData, column: st.delta_generator.DeltaGenerator):
        """Renders the waveform preview (button-driven) and settings controls."""
        with column:
            try:
                original_audio = track_data.get("original_audio")
                track_sr = track_data.get("sr", GLOBAL_SR)
                # Display duration based on original audio if available, else 0
                full_len_samples = len(original_audio) if original_audio is not None else 0
                full_len_sec = full_len_samples / track_sr if track_sr > 0 else 0
                st.caption(f"SR: {track_sr} Hz | Full Duration: {full_len_sec:.2f}s")

                # --- Waveform Visualization (Button-Driven with Caching) ---
                st.markdown(f"**Preview Waveform (First {PREVIEW_DURATION_S}s with Effects)**")
                display_path = None
                preview_displayed_from_cache = False

                # Check if original audio exists before attempting cache check/display
                if original_audio is not None and original_audio.size > 0:
                    current_settings_hash = self._calculate_preview_hash(track_data)
                    stored_hash = track_data.get("preview_settings_hash")
                    stored_path = track_data.get("preview_temp_file_path")

                    if stored_hash is not None and current_settings_hash == stored_hash and stored_path and os.path.exists(stored_path):
                        logger.debug(f"Cache hit for track {track_id}. Using existing preview: {stored_path}")
                        display_path = stored_path
                        preview_displayed_from_cache = True
                    elif stored_path and not os.path.exists(stored_path):
                        logger.warning(f"Cached preview file missing for track {track_id}: {stored_path}. Clearing path.")
                        self.app_state.update_track_param(track_id, "preview_temp_file_path", None)
                        self.app_state.update_track_param(track_id, "preview_settings_hash", None)

                    if display_path:
                        ws_options = WaveSurferOptions(
                            height=100, normalize=True, wave_color="#A020F0", progress_color="#800080", cursor_color="#333333", cursor_width=1, bar_width=2, bar_gap=1
                        )
                        update_count = track_data.get("update_counter", 0)  # Counter changes only when button is clicked
                        audix_key = f"audix_{track_id}_{update_count}"
                        logger.debug(f"Displaying preview: '{track_data.get('name', 'N/A')}' key={audix_key} path={display_path}")
                        audix(data=display_path, sample_rate=track_sr, wavesurfer_options=ws_options, key=audix_key)
                    elif not preview_displayed_from_cache:  # Only show prompt if not displayed from cache
                        st.info("Settings changed or preview not generated. Click 'Update Preview' below.")
                else:
                    # Handle case where track exists (from load) but audio data is missing
                    if track_data.get("source_type") == "upload":
                        st.warning(f"Source file '{track_data.get('original_filename', 'N/A')}' needs to be re-uploaded to generate preview.", icon="⚠️")
                    else:
                        st.info("Track has no audio data.")  # Should ideally not happen for generated tracks post-load
                # -------------------------------------------------------------

                st.markdown("---")
                st.markdown("**Track Settings**")
                st.caption("Adjust settings, then click 'Update Preview' below.")

                # --- Add Subliminalize Button for Affirmation Tracks ---
                if track_data.get("track_type") == TRACK_TYPE_AFFIRMATION:
                    if st.button(
                        "⚡ Subliminalize Preset", key=f"subliminalize_{track_id}", help="Quickly set high speed (4x) and low volume (0.05). Click 'Update Preview' after."
                    ):
                        logger.info(f"Subliminalize preset applied to track {track_id}")
                        self.app_state.update_track_param(track_id, "speed_factor", 4.0)
                        self.app_state.update_track_param(track_id, "volume", 0.05)
                        self.app_state.update_track_param(track_id, "preview_settings_hash", None)  # Invalidate cache
                        st.toast("Subliminal preset applied! Click 'Update Preview'.", icon="⚡")
                        st.rerun()  # Rerun to reflect slider changes
                # ------------------------------------------------------

                col_fx1_1, col_fx1_2, col_fx1_3 = st.columns([0.6, 1, 1])  # Loop/Reverse, Speed, Pitch
                with col_fx1_1:
                    st.markdown("<br/>", unsafe_allow_html=True)
                    loop_value = st.checkbox(
                        "🔁 Loop", key=f"loop_{track_id}", value=track_data.get("loop_to_fit", False), help="Loop track to fit project during final mix/export? (Preview not shown)"
                    )
                    if loop_value != track_data.get("loop_to_fit"):
                        self.app_state.update_track_param(track_id, "loop_to_fit", loop_value)
                    st.markdown("<br/>", unsafe_allow_html=True)  # Add space
                    reverse_value = st.checkbox(
                        "🔄 Reverse", key=f"reverse_{track_id}", value=track_data.get("reverse_audio", False), help="Reverse audio playback? (Updates preview)"
                    )
                    if reverse_value != track_data.get("reverse_audio"):
                        self.app_state.update_track_param(track_id, "reverse_audio", reverse_value)
                with col_fx1_2:
                    speed = st.slider("Speed", 0.25, 4.0, track_data.get("speed_factor", 1.0), 0.05, key=f"speed_{track_id}", help="Playback speed (>1 faster, <1 slower).")
                    if speed != track_data.get("speed_factor"):
                        self.app_state.update_track_param(track_id, "speed_factor", speed)
                with col_fx1_3:
                    pitch = st.slider("Pitch (semitones)", -12, 12, track_data.get("pitch_shift", 0), 1, key=f"pitch_{track_id}", help="Adjust pitch without changing speed.")
                    if pitch != track_data.get("pitch_shift"):
                        self.app_state.update_track_param(track_id, "pitch_shift", pitch)

                col_fx2_1, col_fx2_2, col_fx2_3, col_fx2_4 = st.columns([1, 1, 1, 1])  # Filter, Cutoff, Volume, Pan
                with col_fx2_1:
                    f_type = st.selectbox(
                        "Filter",
                        ["off", "lowpass", "highpass"],
                        index=["off", "lowpass", "highpass"].index(track_data.get("filter_type", "off")),
                        key=f"filter_type_{track_id}",
                        help="Apply low/high pass filter.",
                    )
                    if f_type != track_data.get("filter_type"):
                        self.app_state.update_track_param(track_id, "filter_type", f_type)
                with col_fx2_2:
                    f_enabled = track_data["filter_type"] != "off"
                    max_cutoff = track_sr / 2 - 1
                    f_cutoff = st.number_input(
                        f"Cutoff ({'Hz' if f_enabled else 'Off'})",
                        20.0,
                        max_cutoff if max_cutoff > 20 else 20.0,
                        float(track_data.get("filter_cutoff", 8000.0)),
                        100.0,
                        key=f"filter_cutoff_{track_id}",
                        disabled=not f_enabled,
                        help="Filter cutoff frequency.",
                    )
                    if f_cutoff != track_data.get("filter_cutoff"):
                        self.app_state.update_track_param(track_id, "filter_cutoff", f_cutoff)
                with col_fx2_3:
                    vol = st.slider("Volume", 0.0, 2.0, track_data.get("volume", 1.0), 0.05, key=f"vol_{track_id}", help="Adjust loudness (affects preview after Update).")
                    if vol != track_data.get("volume"):
                        self.app_state.update_track_param(track_id, "volume", vol)
                with col_fx2_4:
                    pan = st.slider("Pan", -1.0, 1.0, track_data.get("pan", 0.0), 0.1, key=f"pan_{track_id}", help="Adjust L/R balance (affects preview after Update).")
                    if pan != track_data.get("pan"):
                        self.app_state.update_track_param(track_id, "pan", pan)

                # --- Update Preview Button ---
                # Disable button if original audio is missing (e.g., loaded project needs re-upload)
                update_disabled = original_audio is None or original_audio.size == 0
                if st.button(
                    "⚙️ Update Preview", key=f"update_preview_{track_id}", help="Generate the 60s preview waveform/audio with the current settings.", disabled=update_disabled
                ):
                    logger.info(f"Update Preview clicked for: '{track_data.get('name', 'N/A')}' ({track_id})")
                    with st.spinner("Generating preview..."):
                        # Calculate hash *before* generating audio
                        new_settings_hash = self._calculate_preview_hash(track_data)
                        preview_audio = get_preview_audio(track_data, preview_duration_s=PREVIEW_DURATION_S)
                        if preview_audio is not None and preview_audio.size > 0:
                            new_preview_path = save_audio_to_temp(preview_audio, track_sr)
                            if new_preview_path:
                                old_preview_path = track_data.get("preview_temp_file_path")
                                # Update state with new path AND hash
                                self.app_state.update_track_param(track_id, "preview_temp_file_path", new_preview_path)
                                self.app_state.update_track_param(track_id, "preview_settings_hash", new_settings_hash)
                                self.app_state.increment_update_counter(track_id)  # Force audix refresh
                                # Clean up old file *after* state update
                                if old_preview_path and old_preview_path != new_preview_path and os.path.exists(old_preview_path):
                                    try:
                                        os.remove(old_preview_path)
                                        logger.info(f"Deleted old preview file: {old_preview_path}")
                                    except OSError as e:
                                        logger.warning(f"Could not delete old preview file {old_preview_path}: {e}")
                                st.toast("Preview updated.", icon="✅")
                            else:
                                st.error("Failed to save new preview file.")
                        else:
                            st.error("Failed to generate preview audio.")
                    st.rerun()  # Rerun to display the new preview
            except Exception as e:
                logger.exception(f"Error rendering main col for {track_id}")
                st.error(f"Error waveform/effects: {e}")

    def _render_track_controls_col(self, track_id: TrackID, track_data: TrackData, column: st.delta_generator.DeltaGenerator) -> bool:
        """Renders the controls (name, type, mute, solo, delete)."""
        # (Implementation remains the same as previous version)
        delete_clicked = False
        with column:
            try:
                st.markdown("**Track Details**")
                name = st.text_input("Name", value=track_data.get("name", "Unnamed"), key=f"name_{track_id}", help="Rename track.")
                if name != track_data.get("name"):
                    self.app_state.update_track_param(track_id, "name", name)
                current_type = track_data.get("track_type", TRACK_TYPE_OTHER)
                try:
                    current_index = TRACK_TYPES.index(current_type)
                except ValueError:
                    current_index = TRACK_TYPES.index(TRACK_TYPE_OTHER)
                new_type = st.selectbox("Type", TRACK_TYPES, index=current_index, key=f"type_{track_id}", help="Categorize this layer (affects icon).")
                if new_type != current_type:
                    self.app_state.update_track_param(track_id, "track_type", new_type)
                    st.rerun()
                st.caption("Mixing (Live Effect)")
                ms_col1, ms_col2 = st.columns(2)
                mute = ms_col1.checkbox("Mute", value=track_data.get("mute", False), key=f"mute_{track_id}", help="Silence track in final mix.")
                if mute != track_data.get("mute"):
                    self.app_state.update_track_param(track_id, "mute", mute)
                solo = ms_col2.checkbox("Solo", value=track_data.get("solo", False), key=f"solo_{track_id}", help="Isolate track(s) in final mix.")
                if solo != track_data.get("solo"):
                    self.app_state.update_track_param(track_id, "solo", solo)
                st.markdown("---")
                if st.button("🗑️ Delete Track", key=f"delete_{track_id}", help="Permanently delete."):
                    delete_clicked = True
                    st.warning(f"Track marked for deletion.")
            except Exception as e:
                logger.exception(f"Error rendering controls for {track_id}")
                st.error(f"Error controls: {e}")
        return delete_clicked

    def render_master_controls(self):
        """Renders the master preview, export, and fade controls."""
        # (Implementation remains the same as v4.3 - Fades removed)
        st.divider()
        st.header("🔊 Master Output")
        # --- REMOVED Fade Controls ---
        st.markdown("---")
        # ---------------------------
        default_filename = "mindmorph_mix"
        export_filename_input = st.text_input(
            "Export Filename (no extension):",
            value=st.session_state.get("export_filename", default_filename),
            key="export_filename",
            help="Enter the desired name for the downloaded file.",
        )
        sanitized_filename = re.sub(r'[\\/*?:"<>|]', "", export_filename_input).strip()
        if not sanitized_filename:
            sanitized_filename = default_filename
        # Store sanitized name in session state for the download button to use
        st.session_state.export_filename_sanitized = sanitized_filename

        if "calculated_mix_duration_s" in st.session_state and st.session_state.calculated_mix_duration_s is not None:
            duration_str = f"{st.session_state.calculated_mix_duration_s:.2f}"
            st.info(f"Estimated Full Mix Duration: **{duration_str} seconds**")
        st.caption("Note: Export time depends on total duration and effects used.")

        master_cols = st.columns(2)
        with master_cols[0]:
            st.button(
                "🎧 Preview Mix (10s)",
                key="preview_mix",
                use_container_width=True,
                help="Generate and play first 10s of the final mix with all effects applied.",
                on_click=self._handle_preview_click,
            )
        with master_cols[1]:
            export_format = st.radio(
                "Export Format:",
                ["WAV", "MP3"],
                key="export_format",
                horizontal=True,
                help="Choose WAV (lossless, large) or MP3 (compressed, smaller)." if PYDUB_AVAILABLE else "WAV format only (MP3 requires pydub/ffmpeg).",
            )
            export_disabled = export_format == "MP3" and not PYDUB_AVAILABLE
            export_button_label = f"💾 Export Full Mix (.{export_format.lower()})"
            export_help = f"Generate the complete final mix as .{export_format.lower()}."
            if export_disabled:
                export_help += " MP3 export disabled (pydub/ffmpeg missing)."
            st.button(export_button_label, key="export_mix", use_container_width=True, help=export_help, on_click=self._handle_export_click, disabled=export_disabled)

            if "export_buffer" in st.session_state and st.session_state.export_buffer:
                file_ext = st.session_state.get("export_file_ext", "wav")
                download_filename = f"{st.session_state.get('export_filename_sanitized', 'mindmorph_mix')}.{file_ext}"  # Use sanitized name
                mime_type = f"audio/{file_ext}"
                st.download_button(
                    label=f"⬇️ Download: {download_filename}",
                    data=st.session_state.export_buffer,
                    file_name=download_filename,
                    mime=mime_type,
                    key="download_export_key",
                    use_container_width=True,
                )
                st.session_state.export_buffer = None
                st.session_state.export_file_ext = None
                st.session_state.calculated_mix_duration_s = None
                st.session_state.export_filename_sanitized = None

    def _handle_preview_click(self):
        """Callback for Preview Mix button."""
        logger.info("Preview Mix button clicked.")
        tracks = self.app_state.get_all_tracks()
        if "preview_audio" in st.session_state:
            del st.session_state.preview_audio
        if "calculated_mix_duration_s" in st.session_state:
            del st.session_state.calculated_mix_duration_s
        if not tracks:
            st.warning("No tracks loaded.")
            return
        with st.spinner("Generating preview mix..."):
            mix_preview, _ = mix_tracks(tracks, preview=True)  # Removed fade args
            if mix_preview is not None and mix_preview.size > 0:
                st.session_state.preview_audio = save_audio(mix_preview, GLOBAL_SR)
                logger.info("Preview generated.")
            else:
                logger.warning("Preview mix empty.")

    def _handle_export_click(self):
        """Callback for Export Mix button. Calculates duration first."""
        logger.info("Export Full Mix button clicked.")
        tracks = self.app_state.get_all_tracks()
        export_format = st.session_state.get("export_format", "WAV").lower()
        # --- Removed fade reads ---
        export_filename_base = st.session_state.get("export_filename", "mindmorph_mix")  # Read from text input state
        # Store final sanitized name for download button
        st.session_state.export_filename_final = export_filename_base

        if "export_buffer" in st.session_state:
            del st.session_state.export_buffer
        if "export_file_ext" in st.session_state:
            del st.session_state.export_file_ext
        if "calculated_mix_duration_s" in st.session_state:
            del st.session_state.calculated_mix_duration_s  # Clear old duration
        if not tracks:
            st.warning("No tracks loaded.")
            return

        # --- Calculate Duration First (Estimate) ---
        calculated_max_len = 0
        processed_lengths = {}
        valid_ids_for_len_calc = []
        solo_active = any(t.get("solo", False) for t in tracks.values())
        with st.spinner("Calculating final mix duration..."):
            for track_id, t_data in tracks.items():
                is_active = t_data.get("solo", False) if solo_active else not t_data.get("mute", False)
                original_audio = t_data.get("original_audio")
                if is_active and original_audio is not None and original_audio.size > 0:
                    original_len = len(original_audio)
                    speed_factor = t_data.get("speed_factor", 1.0)
                    estimated_len = int(original_len / speed_factor) if speed_factor > 0 else original_len
                    processed_lengths[track_id] = estimated_len
                    valid_ids_for_len_calc.append(track_id)
            if valid_ids_for_len_calc:
                calculated_max_len = max(processed_lengths.values())
            st.session_state.calculated_mix_duration_s = calculated_max_len / GLOBAL_SR if GLOBAL_SR > 0 else 0
            logger.info(f"Calculated mix duration (pre-looping, based on speed): {st.session_state.calculated_mix_duration_s:.2f}s")
        # -----------------------------

        # Now generate the actual mix using the sequential method
        with st.spinner(f"Generating full mix ({export_format.upper()})... This may take time."):
            full_mix, final_mix_len_samples = mix_tracks(tracks, preview=False)  # Removed fade args
            if final_mix_len_samples is not None:
                st.session_state.calculated_mix_duration_s = final_mix_len_samples / GLOBAL_SR if GLOBAL_SR > 0 else 0
                logger.info(f"Actual final mix duration (post-looping): {st.session_state.calculated_mix_duration_s:.2f}s")

            if full_mix is not None and full_mix.size > 0:
                if export_format == "wav":
                    st.session_state.export_buffer = save_audio(full_mix, GLOBAL_SR)
                    st.session_state.export_file_ext = "wav"
                    logger.info("Full WAV mix generated.")
                elif export_format == "mp3" and PYDUB_AVAILABLE:
                    try:
                        logger.info("Converting full mix to MP3...")
                        audio_int16 = (full_mix * 32767).astype(np.int16)
                        segment = AudioSegment(audio_int16.tobytes(), frame_rate=GLOBAL_SR, sample_width=audio_int16.dtype.itemsize, channels=2)
                        mp3_buffer = BytesIO()
                        segment.export(mp3_buffer, format="mp3", bitrate="192k")
                        mp3_buffer.seek(0)
                        st.session_state.export_buffer = mp3_buffer
                        st.session_state.export_file_ext = "mp3"
                        logger.info("Full MP3 mix generated.")
                    except Exception as e_mp3:
                        logger.exception("Failed to export as MP3.")
                        st.error(f"MP3 Export Failed: {e_mp3}. Check if ffmpeg is installed and accessible.")
                        st.session_state.export_buffer = None
                        st.session_state.export_file_ext = None
                else:
                    logger.error(f"Unsupported export format requested or pydub missing: {export_format}")
                    st.error(f"Cannot export as {export_format.upper()}.")
                    st.session_state.export_buffer = None
                    st.session_state.export_file_ext = None
            else:
                logger.warning("Full mix empty.")
                st.warning("Generated mix is empty.")
                st.session_state.export_buffer = None
                st.session_state.export_file_ext = None

    def render_preview_audio_player(self):
        """Displays the preview audio player if preview data exists in state."""
        # (Implementation remains the same)
        if "preview_audio" in st.session_state and st.session_state.preview_audio:
            st.markdown("**Mix Preview (10s):**")  # Clarified label
            st.audio(st.session_state.preview_audio, format="audio/wav")
            st.session_state.preview_audio = None

    def render_instructions(self):
        """Renders the instructions expander, expanded if no tracks exist."""
        # --- Updated Instructions (Removed Fades) ---
        tracks_exist = bool(self.app_state.get_all_tracks())
        st.divider()
        with st.expander("📖 Show Instructions & Notes", expanded=not tracks_exist):
            st.markdown("""
              **Welcome to MindMorph!** Create custom subliminal audio by layering affirmations, sounds, & frequencies.

              **Workflow:**

              1.  **➕ Add Tracks (Sidebar):**
                  * Use the options on the left to add your audio layers. The *full* audio is loaded/generated.

              2.  **🎚️ Edit Tracks (Main Panel):**
                  * Each track appears below. Set its **Type** (Affirmation, Background, etc.) using the dropdown.
                  * **Track Settings (Reverse, Speed, Pitch, Filter, Loop, Volume, Pan):** Adjust these settings.
                  * **Click `⚙️ Update Preview`** to generate a 60-second preview waveform/audio incorporating **all** current settings for that track. *(Note: Looping effect is only calculated during final mix/export, not shown in preview).*
                  * **Track Controls (Mute, Solo):** These affect the final mix directly without needing 'Update Preview'.
                  * `Track Name`: Rename tracks.
                  * `🗑️ Delete Track`: Remove unwanted tracks.

              3.  **🔊 Mix & Export (Bottom Panel):**
                  * Enter a **Filename** for your download.
                  * `🎧 Preview Mix (10s)`: Hear the start of the combined audio. This processes the *full duration* of all tracks with *all* saved settings.
                  * `💾 Export Full Mix`: Choose WAV or MP3 (if available). This generates the final file using the *full* audio tracks and all settings. Click the download button.

              **Tips:**
              * The editor preview is limited to 60s for performance, but the final export uses the full audio.
              * Use low volume for subliminal affirmation tracks (adjust and click 'Update Preview' to hear).
              * High speed (e.g., 2x-4x) is common for affirmation tracks.
              * Use the `🔁 Loop` option for short sounds; it will be applied during Mix/Export.
              """)


# ==========================================
# 5. Benchmarking (Placeholder Class)
# ==========================================
# (Benchmarker class remains the same)
class Benchmarker:
    """Placeholder for performance benchmarking."""

    def __init__(self, tts_generator: TTSGenerator):
        self.tts_generator = tts_generator
        logger.debug("Benchmarker initialized.")

    def benchmark_tts(self, word_count: int = 10000, repetitions: int = 1):
        logger.info(f"Starting TTS benchmark: {word_count} words, {repetitions} reps.")
        st.subheader(f"Benchmark: TTS Generation ({word_count} words)")
        if word_count <= 0:
            st.warning("Word count must be positive.")
            return
        placeholder_word = "benchmark "
        text = (placeholder_word * (word_count // len(placeholder_word) + 1))[: word_count * 6]
        st.text(f"Generating ~{word_count} words...")
        total_time = 0
        min_time = float("inf")
        max_time = 0
        success_count = 0
        progress_bar = st.progress(0)
        status_text = st.empty()
        for i in range(repetitions):
            status_text.text(f"Repetition {i + 1}/{repetitions}...")
            start_time = time.time()
            try:
                audio, sr = self.tts_generator.generate(text)
                end_time = time.time()
                if audio is not None and audio.size > 0:
                    duration = end_time - start_time
                    total_time += duration
                    min_time = min(min_time, duration)
                    max_time = max(max_time, duration)
                    success_count += 1
                    logger.info(f"TTS Benchmark rep {i + 1} success in {duration:.2f}s")
                else:
                    logger.error(f"TTS Benchmark rep {i + 1} failed (no audio).")
                    st.error(f"Rep {i + 1} failed.")
            except Exception as e:
                logger.exception(f"TTS Benchmark rep {i + 1} failed.")
                st.error(f"Rep {i + 1} failed: {e}")
            progress_bar.progress((i + 1) / repetitions)
        status_text.text("Benchmark complete.")
        if success_count > 0:
            avg_time = total_time / success_count
            st.metric("Average Time per Generation", f"{avg_time:.2f} s")
            st.text(f"Min: {min_time:.2f} s | Max: {max_time:.2f} s | Success: {success_count}/{repetitions}")
        else:
            st.error("TTS Benchmark failed for all repetitions.")


# ==========================================
# 6. Main Application Logic & Onboarding V5
# ==========================================
# (Onboarding Welcome Message remains the same as v2.4)
def show_welcome_message():
    """Displays the welcome message container using columns for clarity."""
    if "welcome_message_shown" not in st.session_state:
        with st.container(border=True):
            st.markdown("### 👋 Welcome to MindMorph!")
            st.markdown("Create custom subliminal audio by layering sounds and applying effects.")
            st.markdown("---")
            st.markdown("#### Quick Start:")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown("#### 1. Add Tracks ➕")
                st.markdown("Look Left! **👈** Use the **sidebar**.")
                st.caption("Upload, Generate TTS, Tones")
            with col2:
                st.markdown("#### 2. Edit Tracks 🎚️")
                st.markdown("Adjust **settings** in the main panel.")
                st.caption("Click 'Update Preview' to refresh!")  # Updated caption
            with col3:
                st.markdown("#### 3. Mix & Export 🔊")
                st.markdown("Use **master controls** at the bottom.")
                st.caption("Preview or Download WAV")
            st.markdown("---")
            st.markdown("*(Click button below to hide this guide. Find details in Instructions at page bottom.)*")
            button_cols = st.columns([1, 1.5, 1])  # Centering columns
            with button_cols[1]:
                if st.button("Got it! Let's Start Creating ✨", key="dismiss_welcome", type="primary", use_container_width=True):
                    st.session_state.welcome_message_shown = True
                    logger.info("Welcome message dismissed.")
                    st.rerun()


def main():
    """Main function to run the Streamlit application."""
    logger.info("Starting main application function.")
    st.title("🧠 MindMorph - Subliminal Audio Editor")
    show_welcome_message()  # Show welcome message if needed

    # --- Project Load Handling ---
    if st.session_state.get("project_load_requested", False):
        logger.info("Project load requested.")
        loaded_data = st.session_state.get("uploaded_project_file_data")
        st.session_state.project_load_requested = False  # Reset flag
        st.session_state.uploaded_project_file_data = None

        if loaded_data:
            try:
                project_content = json.loads(loaded_data.decode("utf-8"))
                # TODO: Add version check project_content.get("version") == PROJECT_FILE_VERSION
                if "tracks" in project_content:
                    logger.info("Valid project structure found. Clearing current state.")
                    # Clear existing state
                    app_state_temp = AppState()  # Get instance to access methods
                    current_tracks = app_state_temp.get_all_tracks()
                    for track_id in list(current_tracks.keys()):
                        app_state_temp.delete_track(track_id)  # Ensures preview files are cleaned

                    logger.info("Loading tracks from file...")
                    loaded_tracks_data = project_content["tracks"]
                    # Need access to TTS generator and other generators here
                    tts_gen_temp = TTSGenerator()
                    ui_manager_temp = UIManager(app_state_temp, tts_gen_temp)  # Need UIManager for generators

                    with st.spinner("Reconstructing project..."):
                        for track_id, track_data in loaded_tracks_data.items():
                            logger.debug(f"Loading track {track_id}: {track_data.get('name')}")
                            source_type = track_data.get("source_type", "unknown")
                            track_type = track_data.get("track_type", TRACK_TYPE_OTHER)
                            new_track_id = str(uuid.uuid4())  # Generate new ID on load

                            # Reconstruct original_audio based on source type
                            reconstructed_audio = None
                            if source_type == "tts" and "tts_text" in track_data:
                                logger.info(f"Regenerating TTS for track {track_id}")
                                reconstructed_audio, _ = tts_gen_temp.generate(track_data["tts_text"])
                            elif source_type == "noise" and "gen_noise_type" in track_data:
                                logger.info(f"Regenerating {track_data['gen_noise_type']} for track {track_id}")
                                reconstructed_audio = generate_noise(track_data["gen_noise_type"], track_data.get("gen_duration", 60), GLOBAL_SR, track_data.get("gen_volume", 0.5))
                            elif source_type == "binaural" and "gen_freq_left" in track_data:
                                logger.info(f"Regenerating Binaural Beats for track {track_id}")
                                reconstructed_audio = generate_binaural_beats(
                                    track_data.get("gen_duration", 60),
                                    track_data["gen_freq_left"],
                                    track_data.get("gen_freq_right", 210),
                                    GLOBAL_SR,
                                    track_data.get("gen_volume", 0.3),
                                )
                            elif source_type == "solfeggio" and "gen_freq" in track_data:
                                logger.info(f"Regenerating Solfeggio for track {track_id}")
                                reconstructed_audio = generate_solfeggio_frequency(
                                    track_data.get("gen_duration", 60), track_data["gen_freq"], GLOBAL_SR, track_data.get("gen_volume", 0.3)
                                )
                            elif source_type == "isochronic" and "gen_carrier_freq" in track_data:
                                logger.info(f"Regenerating Isochronic for track {track_id}")
                                reconstructed_audio = generate_isochronic_tones(
                                    track_data.get("gen_duration", 60),
                                    track_data["gen_carrier_freq"],
                                    track_data.get("gen_pulse_freq", 10),
                                    GLOBAL_SR,
                                    track_data.get("gen_volume", 0.4),
                                )
                            elif source_type == "upload":
                                logger.warning(f"Track {track_id} is an upload ('{track_data.get('original_filename')}'). Audio data not saved in project file.")
                                # Keep original_audio as None - UI will show warning
                            else:
                                logger.warning(f"Unknown or missing source_type for track {track_id}. Cannot reconstruct audio.")

                            # Add track to state (with potentially None audio for uploads)
                            track_data["original_audio"] = reconstructed_audio
                            app_state_temp.add_track(new_track_id, track_data, track_type=track_type)

                    st.success("Project loaded successfully!")
                    # No st.rerun() needed here as we are already at the top of the script execution
                else:
                    st.error("Invalid project file structure.")
            except json.JSONDecodeError:
                st.error("Failed to decode project file. Ensure it's a valid JSON file.")
            except Exception as e:
                logger.exception("Error loading project file.")
                st.error(f"An error occurred while loading the project: {e}")
        else:
            logger.warning("Project load requested but no file data found in session state.")
    # --- End Project Load Handling ---

    if "welcome_message_shown" in st.session_state:
        st.markdown("Create custom subliminals by layering affirmations, sounds, & frequencies. Adjust effects & mix.")
        st.divider()
        app_state = AppState()  # Re-initialize or get existing state
        tts_generator = TTSGenerator()
        ui_manager = UIManager(app_state, tts_generator)

        ui_manager.render_sidebar()
        ui_manager.render_tracks_editor()
        ui_manager.render_master_controls()
        ui_manager.render_preview_audio_player()
        ui_manager.render_instructions()

        st.divider()
        st.caption("MindMorph Subliminal Editor")
        logger.info("Reached end of main application function render.")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        logger.exception("A critical error occurred in main execution.")
        logger.error(f"Unhandled exception in main: {e}")
        st.error("An unexpected error occurred. Please check the application logs or try reloading.")
        # st.code(traceback.format_exc()) # Avoid showing full traceback in UI
