# utils.py
# ==========================================
# General Utility Functions for MindMorph
# ==========================================

import logging
import logging.handlers
import os
import queue
from typing import Optional

# Need to import docx if read_text_file uses it
try:
    import docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

    # Define dummy Document if needed, or handle absence in read_text_file
    class Document:
        pass  # Dummy class


import streamlit as st  # Used for session_state in logging and error display
from streamlit.runtime.uploaded_file_manager import UploadedFile

# Import constants from config
from config import LOG_BACKUP_COUNT, LOG_FILE, LOG_FORMAT, LOG_MAX_BYTES

# Get a logger for this module
logger = logging.getLogger(__name__)

# --- Logging Setup ---
# Define the queue at the module level
log_queue = queue.Queue(-1)


def setup_logging():
    """Configures the application's logging using a queue handler."""
    # Prevent adding handlers multiple times if setup_logging is called again
    # This can happen with Streamlit's rerun behavior
    root_logger = logging.getLogger()
    # Check if a QueueHandler instance already exists
    if any(isinstance(h, logging.handlers.QueueHandler) for h in root_logger.handlers):
        logger.debug("Logging handlers already appear to be configured.")
        return

    logger.info(f"Setting up logging. Log file: {LOG_FILE}")
    log_formatter = logging.Formatter(LOG_FORMAT)

    # File handler (rotates logs) - This will be used by the listener
    try:
        file_handler = logging.handlers.RotatingFileHandler(LOG_FILE, maxBytes=LOG_MAX_BYTES, backupCount=LOG_BACKUP_COUNT, encoding="utf-8")
        file_handler.setFormatter(log_formatter)
        file_handler.setLevel(logging.DEBUG)  # Log DEBUG level and above to file
    except Exception as e:
        st.error(f"Error setting up log file handler: {e}")
        logger.error(f"Error setting up log file handler: {e}")
        return  # Stop setup if file handler fails

    # Queue handler (sends logs from the main thread to the queue)
    queue_handler = logging.handlers.QueueHandler(log_queue)
    queue_handler.setLevel(logging.DEBUG)  # Send all logs to the queue

    # Add ONLY the queue handler to the root logger in the main process
    root_logger.addHandler(queue_handler)
    root_logger.setLevel(logging.DEBUG)  # Set root logger level

    # Start the queue listener thread (if not already started)
    # Use session_state to track if the listener is running
    if "log_listener_started" not in st.session_state:
        try:
            listener = logging.handlers.QueueListener(log_queue, file_handler, respect_handler_level=True)
            listener.start()
            st.session_state.log_listener_started = True
            logger.info("Logging QueueListener started.")
        except Exception as e:
            st.error(f"Error starting log listener: {e}")
            logger.error(f"Error starting log listener: {e}")
            # Attempt to remove the queue handler if listener fails to start
            root_logger.removeHandler(queue_handler)
    else:
        logger.debug("Logging QueueListener already running.")


# --- File Reading ---


def read_text_file(uploaded_file: UploadedFile) -> Optional[str]:
    """
    Reads content from uploaded txt or docx file.

    Handles potential decoding errors for txt and parsing errors for docx.

    Args:
        uploaded_file: The UploadedFile object from Streamlit.

    Returns:
        The text content as a string, or None if reading fails or the
        file type is unsupported.
    """
    if uploaded_file is None:
        logger.error("read_text_file called with None object.")
        return None

    file_name = uploaded_file.name
    file_type = uploaded_file.type
    logger.info(f"Attempting to read file: {file_name}, Type: {file_type}")

    try:
        # --- Read .txt file ---
        if file_name.lower().endswith(".txt") or file_type == "text/plain":
            logger.debug(f"Reading .txt file: {file_name}")
            try:
                # Read bytes and decode with UTF-8, replacing errors
                file_content_bytes = uploaded_file.getvalue()
                return file_content_bytes.decode("utf-8", errors="replace")
            except Exception as e_decode:
                logger.exception(f"Error decoding .txt file {file_name} as UTF-8.")
                # TODO: Remove direct Streamlit call. Raise exception or return status.
                st.error(f"Error decoding text file '{file_name}': {e_decode}. Ensure it's UTF-8 encoded.")
                return None

        # --- Read .docx file ---
        elif file_name.lower().endswith(".docx") or file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            if not DOCX_AVAILABLE:
                logger.error("Attempted to read .docx file, but 'python-docx' library is not installed.")
                st.error("Cannot read .docx files. Please install the 'python-docx' library (`pip install python-docx`).")
                return None

            logger.debug(f"Processing .docx file: {file_name}")
            try:
                # Need to operate on the file-like object directly
                uploaded_file.seek(0)  # Ensure reading from the start
                document = docx.Document(uploaded_file)
                logger.debug(f"docx document parsed. Found {len(document.paragraphs)} paragraphs.")

                # Extract text from paragraphs
                para_texts = []
                for i, para in enumerate(document.paragraphs):
                    if para is None:
                        logger.warning(f"Paragraph {i} in '{file_name}' is None. Skipping.")
                        continue
                    # Ensure paragraph text is not None before stripping/appending
                    para_text = getattr(para, "text", "")
                    if para_text and para_text.strip():  # Only add non-empty paragraphs
                        para_texts.append(para_text.strip())

                logger.debug(f"Extracted {len(para_texts)} non-empty paragraphs from '{file_name}'.")
                return "\n".join(para_texts)  # Join paragraphs with newlines

            except Exception as e_docx:
                logger.exception(f"Error processing .docx file: {file_name}")
                # TODO: Remove direct Streamlit call. Raise exception or return status.
                st.error(f"Error reading Word document '{file_name}': {e_docx}")
                return None

        # --- Unsupported file type ---
        else:
            logger.error(f"Unsupported file type for text reading: {file_type} (Filename: {file_name})")
            # TODO: Remove direct Streamlit call. Raise exception or return status.
            st.error(f"Unsupported file type: '{file_type}'. Please upload a .txt or .docx file.")
            return None

    except Exception as e_outer:
        # Catch any other unexpected errors during file handling
        logger.exception(f"An unexpected error occurred while reading file: {file_name}")
        # TODO: Remove direct Streamlit call. Raise exception or return status.
        st.error(f"Error reading file '{file_name}': {e_outer}")
        return None


# Potential future additions:
# - Function to sanitize filenames
# - etc.
